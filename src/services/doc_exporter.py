"""
Document Exporter — exports tailored resumes to DOCX format.

Uses python-docx to create professional, formatted DOCX files from
structured resume data and tailored text.
"""

from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocExporter:
    """Exports resume data to professionally formatted DOCX files."""

    def __init__(self, output_dir: str | Path | None = None) -> None:
        self._output_dir = Path(
            output_dir
            or os.getenv("OUTPUTS_DIR", "outputs")
        )
        self._output_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        resume_text: str,
        contact: dict[str, Any] | None = None,
        sections: dict[str, list[str]] | None = None,
        filename: str | None = None,
        job_title: str | None = None,
        company: str | None = None,
    ) -> str:
        """
        Export resume text to a DOCX file.

        Args:
            resume_text: The resume text to export
            contact: Contact info dict (name, email, phone)
            sections: Organised sections dict
            filename: Custom filename (auto-generated if None)
            job_title: Job title (for auto-naming)
            company: Company name (for auto-naming)

        Returns:
            Path to the created DOCX file
        """
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # If we have structured sections, use them for formatting
        if sections and contact:
            self._build_structured(doc, contact, sections)
        else:
            self._build_from_text(doc, resume_text)

        # Generate filename
        if not filename:
            filename = self._generate_filename(contact, job_title, company)

        output_path = self._output_dir / filename
        doc.save(str(output_path))

        return str(output_path.resolve())

    def _build_structured(
        self,
        doc: Document,
        contact: dict[str, Any],
        sections: dict[str, list[str]],
    ) -> None:
        """Build a structured DOCX from organised resume data."""
        # Name header
        name = contact.get("name", "")
        if name:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(name)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Contact line
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        if contact.get("github"):
            contact_parts.append(contact["github"])

        if contact_parts:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(" | ".join(contact_parts))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        # Sections: render user-supplied keys first (preserving resume order),
        # then any known-order sections that remain.
        known_order = [
            "summary", "experience", "education", "skills",
            "projects", "certifications", "awards", "publications",
            "languages", "other",
        ]
        seen = set()
        ordered_keys: list[str] = []
        for key in sections:
            ordered_keys.append(key)
            seen.add(key)
        for key in known_order:
            if key not in seen:
                ordered_keys.append(key)

        for section_name in ordered_keys:
            content = sections.get(section_name, [])
            if not content:
                continue

            # Section header
            heading = doc.add_heading(section_name.title(), level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 50, 100)

            # Section content
            for item in content:
                if item.strip():
                    para = doc.add_paragraph(item.strip(), style="List Bullet")
                    para.paragraph_format.space_after = Pt(2)

    def _build_from_text(self, doc: Document, text: str) -> None:
        """Build DOCX from plain text, preserving structure."""
        lines = text.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # Detect section headers (ALL CAPS or short lines)
            if stripped.isupper() and len(stripped) < 50:
                heading = doc.add_heading(stripped.title(), level=2)
                heading.runs[0].font.color.rgb = RGBColor(0, 50, 100)
            elif stripped.startswith(("•", "-", "–", "▪", "►")):
                # Bullet points
                bullet_text = stripped.lstrip("•-–▪► ").strip()
                doc.add_paragraph(bullet_text, style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    def _generate_filename(
        self,
        contact: dict[str, Any] | None,
        job_title: str | None,
        company: str | None,
    ) -> str:
        """Generate a descriptive filename."""
        parts = []

        name = contact.get("name", "") if contact else ""
        if name:
            parts.append(name.replace(" ", "_"))
        else:
            parts.append("resume")

        if company:
            parts.append(company.replace(" ", "_"))
        if job_title:
            parts.append(job_title.replace(" ", "_"))

        timestamp = datetime.now().strftime("%Y%m%d")
        parts.append(timestamp)

        filename = "_".join(parts)
        # Clean filename — keep alphanumeric, underscores, hyphens, dots
        filename = "".join(c for c in filename if c.isalnum() or c in "_-.")
        return f"{filename}.docx"
